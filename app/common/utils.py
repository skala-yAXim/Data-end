import json
import re

from docx import Document
from openpyxl import load_workbook
from typing import List
from langchain.text_splitter import RecursiveCharacterTextSplitter
from datetime import datetime, timezone, timedelta

def clean_html(raw_html):
    return re.sub(r'<[^>]+>', '', raw_html)
  
def extract_text_from_json(json_str):
    def recursive_extract(obj, results):
        if isinstance(obj, dict):
            for key, value in obj.items():
                if key == "text":
                    results.append(value)
                else:
                    recursive_extract(value, results)
        elif isinstance(obj, list):
            for item in obj:
                recursive_extract(item, results)

    try:
        parsed = json.loads(json_str)
        texts = []
        recursive_extract(parsed, texts)
        return texts
    except json.JSONDecodeError as e:
        print("Invalid JSON:", e)
        return []

def extract_from_docx(file_path: str) -> str:
    """DOCX 파일에서 텍스트 추출"""
    doc = Document(file_path)
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())
            
    # 표 내용도 추출
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n".join(text_parts)

def extract_from_xlsx(file_path: str) -> List[str]:
    """XLSX 파일에서 텍스트 추출"""
    wb = load_workbook(file_path, data_only=True)
    texts = []

    for sheet in wb.sheetnames:
        ws = wb[sheet]
        lines = [f"[시트: {sheet}]"]

        for row in ws.iter_rows():
            row_cells = []
            for cell in row:
                val = cell.value
                text = str(val) if val is not None else ""
                row_cells.append(text)
            
            # 빈 행은 제외
            if any(cell.strip() for cell in row_cells):
                lines.append(" | ".join(row_cells))

        texts.append("\n".join(lines))

    return texts


def extract_from_txt(file_path: str) -> str:
    """TXT 파일에서 텍스트 추출"""
    encodings = ['utf-8', 'cp949', 'euc-kr']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    
    # 모든 인코딩 실패 시
    return "텍스트 파일 인코딩 오류"

def split_into_chunks(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[str]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_text(text)

def convert_utc_to_kst(utc_datetime_str: str) -> datetime:
    """
    UTC ISO 8601 문자열(밀리초 포함 가능)을 받아 한국 시간(KST) datetime 객체로 반환.
    빈 문자열이나 파싱 실패 시 현재 KST 시간을 반환.
    """
    try:
        if not utc_datetime_str:
            raise ValueError("빈 문자열")
        
        # 밀리초(.f)를 처리할 수 있도록 포맷을 수정하고, 끝의 Z를 제거
        if utc_datetime_str.endswith('Z'):
            utc_datetime_str = utc_datetime_str[:-1]

        # 마이크로초가 없는 경우와 있는 경우 모두 처리
        try:
            # 밀리초가 있는 경우
            utc_time = datetime.strptime(utc_datetime_str, "%Y-%m-%dT%H:%M:%S.%f")
        except ValueError:
            # 밀리초가 없는 경우
            utc_time = datetime.strptime(utc_datetime_str, "%Y-%m-%dT%H:%M:%S")

        # 타임존 정보(UTC)를 설정
        utc_time = utc_time.replace(tzinfo=timezone.utc)

    except Exception:
        # 파싱 실패 시 (예: 포맷이 아예 다른 경우) 현재 시간으로 대체
        utc_time = datetime.now(timezone.utc)

    kst = timezone(timedelta(hours=9))
    return utc_time.astimezone(kst)
