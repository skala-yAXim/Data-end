from datetime import datetime, timezone
import os
from tempfile import mkdtemp
import tempfile
from typing import List, Optional
from app.common.config import MICROSOFT_CLIENT_ID, MICROSOFT_CLIENT_SECRET, MICROSOFT_TENANT_ID
from app.schemas.docs_activity import DocsEntry
from msal import ConfidentialClientApplication
import requests
from dateutil.parser import parse
import pandas as pd
from docx import Document

from app.vectordb.schema import BaseRecord, DocumentMetadata
from app.vectordb.uploader import upload_data_to_db

def get_access_token(client_id: str, client_secret: str, tenant_id: str):
    # Graph API 설정
    authority = f"https://login.microsoftonline.com/{tenant_id}"
    scope = ["https://graph.microsoft.com/.default"]

    # MSAL 앱 초기화
    app = ConfidentialClientApplication(
        client_id,
        authority=authority,
        client_credential=client_secret
    )

    # 액세스 토큰 요청
    result = app.acquire_token_for_client(scopes=scope)

    if "access_token" in result:
        return result["access_token"]
    else:
        # 에러 메시지 출력
        error = result.get("error", "unknown_error")
        error_description = result.get("error_description", "No description provided.")
        raise Exception(f"토큰 요청 실패: {error} - {error_description}")

def get_sharepoint_domain(access_token: str) -> str:
    url = "https://graph.microsoft.com/v1.0/sites/root"
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(url, headers=headers)

    if response.status_code == 200:
        web_url = response.json().get("webUrl", "")
        domain = web_url.replace("https://", "").split("/")[0]
        return domain
    else:
        raise Exception(f"SharePoint 도메인 조회 실패: {response.status_code} - {response.text}")

def get_drive_id(access_token: str, site_id: str) -> str:
    url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive"
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        return response.json().get("id")
    else:
        raise Exception(f"드라이브 ID 조회 실패: {response.status_code} - {response.text}")

def fetch_drive_files(access_token: str, drive_id: str, folder_id: Optional[str] = None) -> List[DocsEntry]:
    entries: List[DocsEntry] = []

    # 폴더 경로 설정
    if folder_id:
        url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{folder_id}/children"
    else:
        url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root/children"

    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(url, headers=headers)

    if response.status_code != 200:
        raise Exception(f"파일 목록 조회 실패: {response.status_code} - {response.text}")

    data = response.json().get("value", [])

    for item in data:
        filename = item.get("name")
        size = item.get("size", 0)
        last_modified = item.get("lastModifiedDateTime")
        url_link = item.get("webUrl", "")
        file_type = (
            item.get("file", {}).get("mimeType") if "file" in item
            else "folder" if "folder" in item else "unknown"
        )
        file_id = item.get("id", "unknown")


        authors = set()

        # 작성자 정보
        created_by = item.get("createdBy", {}).get("user", {})
        if created_by:
            email = created_by.get("email") or created_by.get("displayName", "알 수 없음")
            authors.add(email)

        # 파일 버전 기록 확인
        if "file" in item:
            versions_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item['id']}/versions"
            versions_response = requests.get(versions_url, headers=headers)

            if versions_response.status_code == 200:
                versions = versions_response.json().get("value", [])
                for version in versions:
                    modified_by = version.get("lastModifiedBy", {}).get("user", {})
                    if modified_by:
                        email = modified_by.get("email") or modified_by.get("displayName", "알 수 없음")
                        authors.add(email)

        # 폴더면 재귀적으로 내부 파일 가져오기
        if "folder" in item:
            folder_items = fetch_drive_files(access_token, drive_id, item["id"])
            entries.extend(folder_items)
        else:
            entry = DocsEntry(
                filename=filename,
                author=list(authors),
                last_modified=datetime.fromisoformat(last_modified.replace("Z", "+00:00")) if last_modified else None,
                type=file_type,
                size=size,
                file_id=file_id,
                drive_id=drive_id
            )
            entries.append(entry)

    return entries

def fetch_all_sites_files(access_token: str, domain: str) -> List[DocsEntry]:
    # 1. 전체 사이트 목록 가져오기
    list_sites_url = "https://graph.microsoft.com/v1.0/sites?search=*"
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(list_sites_url, headers=headers)

    if response.status_code != 200:
        raise Exception(f"사이트 목록 조회 실패: {response.status_code} - {response.text}")

    sites = response.json().get("value", [])
    all_entries: List[DocsEntry] = []

    for site in sites:
        site_id = site.get("id")
        site_name = site.get("name")
        site_url = site.get("webUrl", "")
        
        if not site_id:
            print(f"[건너뜀] site_id 없음: {site}")
            continue
        
        print(f"[시도 중] 사이트 이름: {site_name}, 주소: {site_url}")

        try:
            # 2. 사이트 ID 조회
            site_id = site_id
            if not site_id:
                continue

            # 3. 드라이브 ID 조회
            drive_id = get_drive_id(access_token, site_id)
            if not drive_id:
                continue

            # 4. 드라이브 파일 목록 조회
            entries = fetch_drive_files(access_token, drive_id)
            all_entries.extend(entries)

        except Exception as e:
            print(f"[오류] 사이트 {site_name} 처리 중 오류 발생: {str(e)}")
            continue

    return all_entries

async def save_docs_data():
    # TODO: 오늘 날짜 데이터만 긁어올 수 있도록 수정
    token = get_access_token(client_id=MICROSOFT_CLIENT_ID, client_secret=MICROSOFT_CLIENT_SECRET, tenant_id=MICROSOFT_TENANT_ID)
    domain = get_sharepoint_domain(token)
    
    docs_data = fetch_all_sites_files(token, domain)
    
    records = []

    for entry in docs_data:
        content = extract_file_content(entry, token)
        if content is None:
            content = ""  # 빈 문자열로 대체하거나 continue 할 수도 있음
        
        record = create_record_from_entry(content, entry)
        records.append(record)
    
    collection_name = "Documents"
    
    upload_data_to_db(collection_name=collection_name, records=records)
        
    return docs_data

def extract_file_content(docs_entry: DocsEntry, access_token: str) -> Optional[str]:
    
    tmp_dir = tempfile.mkdtemp(prefix="docs_dl_")
    file_path = os.path.join(tmp_dir, docs_entry.filename)
    headers = {"Authorization": f"Bearer {access_token}"}

    download_url = f"https://graph.microsoft.com/v1.0/drives/{docs_entry.drive_id}/items/{docs_entry.file_id}/content"

    try:
        response = requests.get(download_url, headers=headers)
        if response.status_code != 200:
            return f"파일 다운로드 실패: {response.status_code} - {response.text}"

        with open(file_path, "wb") as f:
            f.write(response.content)

    except Exception as e:
        return f"파일 다운로드 중 오류: {str(e)}"
    
    # 파일 형식에 따라 추출
    file_ext = docs_entry.filename.lower().split('.')[-1]
    try:
        if file_ext == 'docx':
            return _extract_from_docx(file_path)
        elif file_ext == 'xlsx':
            return _extract_from_xlsx(file_path)
        elif file_ext == 'txt':
            return _extract_from_txt(file_path)
        else:
            print("지원하지 않는 파일 형식")
            return f"지원하지 않는 파일 형식: {file_ext}"
    except Exception as e:
        return f"파일 읽기 오류: {str(e)}"
    finally:
        # 파일 정리 필요시 삭제 가능
        try:
            os.remove(file_path)
            os.rmdir(tmp_dir)
        except Exception:
            pass

def _extract_from_docx(file_path: str) -> str:
    """DOCX 파일에서 텍스트 추출"""
    doc = Document(file_path)
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())
            
    # 표 내용도 추출
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n".join(text_parts)

def _extract_from_xlsx(file_path: str) -> str:
    """XLSX 파일에서 텍스트 추출"""
    text_parts = []
    
    # 모든 시트 읽기
    excel_file = pd.ExcelFile(file_path)
    for sheet_name in excel_file.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        
        text_parts.append(f"[시트: {sheet_name}]")
        
        # 컬럼명 추가
        if not df.empty:
            text_parts.append("컬럼: " + " | ".join(str(col) for col in df.columns))
            
            # 데이터 행 추가 (최대 100행)
            for idx, row in df.head(100).iterrows():
                row_text = " | ".join(str(val) for val in row.values if pd.notna(val))
                if row_text.strip():
                    text_parts.append(row_text)
    
    return "\n".join(text_parts)

def _extract_from_txt(file_path: str) -> str:
    """TXT 파일에서 텍스트 추출"""
    encodings = ['utf-8', 'cp949', 'euc-kr']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    
    # 모든 인코딩 실패 시
    return "텍스트 파일 인코딩 오류"

def create_record_from_entry(content: str, entry: DocsEntry) -> BaseRecord[DocumentMetadata]:
    return BaseRecord[DocumentMetadata](
        text=content.strip(),
        metadata=DocumentMetadata(
            file_id=entry.file_id,
            filename=entry.filename,
            author=entry.author,
            last_modified=entry.last_modified,
            type=entry.type,
            size=entry.size,
        )
    )
